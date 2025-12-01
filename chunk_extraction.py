
import re
from docx import Document
from typing import List, Dict


def extract_chunks(docx_path: str) -> List[Dict]:
	"""
	Extracts chunks from a .docx file so that EVERY numbered section (e.g., 1, 1.1, 1.2, 2, 2.1, etc.) becomes a chunk, even if it's just a heading/title.
	Returns a list of dicts: {id, title, content, path, source}
	"""
	import os
	doc = Document(docx_path)
	# DEBUG: Print all paragraph texts with their indexes for regex tuning
	print("--- Paragraphs in document ---")
	for idx, para in enumerate(doc.paragraphs):
		print(f"{idx}: [{para.text.strip()}]")
	print("--- End of paragraphs ---\n")
	chunks = []
	current_chunk = None
	chunk_id = 0
	# Match lines starting with '#'
	section_marker_re = re.compile(r'^#\s*(.*)')
	source = os.path.splitext(os.path.basename(docx_path))[0]
	for para in doc.paragraphs:
		text = para.text.strip()
		if not text:
			continue
		match = section_marker_re.match(text)
		if match:
			# Save previous chunk
			if current_chunk:
				chunks.append(current_chunk)
			chunk_id += 1
			title = match.group(1).strip() or text
			current_chunk = {
				'id': str(chunk_id),
				'title': title,
				'content': '',
				'path': [str(chunk_id)],
				'source': source
			}
		elif current_chunk:
			# Add content to current chunk
			if current_chunk['content']:
				current_chunk['content'] += '\n'
			current_chunk['content'] += text
	# Add last chunk
	if current_chunk:
		chunks.append(current_chunk)
	return chunks

	# Add last chunk
	if current_chunk:
		chunks.append(current_chunk)

	return chunks

# Usage: python chunk_extraction.py <docx_path>
if __name__ == "__main__":
	import sys, json, os
	if len(sys.argv) < 2:
		print("Usage: python chunk_extraction.py <docx_path>")
		exit(1)
	docx_path = sys.argv[1]
	# Extract chunks and save to a file named after the docx
	chunks = extract_chunks(docx_path)
	base = os.path.splitext(os.path.basename(docx_path))[0]
	out_path = f"{base}_chunks.json"
	with open(out_path, "w", encoding="utf-8") as f:
		json.dump(chunks, f, indent=2, ensure_ascii=False)
	print(f"Extracted {len(chunks)} chunks and saved to {out_path}")
# chunk_extraction.py
# Extracts hierarchical chunks from .docx with metadata
